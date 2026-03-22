from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent


def set_rtl(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def style_run(run, font_name="Times New Roman", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def apply_doc_defaults(document: Document, rtl: bool):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    if rtl:
        for section in document.sections:
            sect_pr = section._sectPr
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            sect_pr.append(bidi)


def add_heading(document: Document, text: str, level: int, rtl: bool):
    paragraph = document.add_heading(level=min(level, 4))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    style_run(run, size=max(14, 20 - (level * 2)), bold=True)
    if rtl:
        set_rtl(paragraph)


def add_paragraph(document: Document, text: str, rtl: bool, style_name: str | None = None):
    paragraph = document.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    style_run(run)
    if rtl:
        set_rtl(paragraph)
    return paragraph


def add_bullet(document: Document, text: str, rtl: bool, ordered: bool = False):
    style_name = "List Number" if ordered else "List Bullet"
    return add_paragraph(document, text, rtl, style_name=style_name)


def add_image(document: Document, image_path: Path, caption: str | None, rtl: bool):
    if not image_path.exists():
        add_paragraph(document, f"[Missing image: {image_path.name}]", rtl)
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if rtl:
        set_rtl(paragraph)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    if caption:
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        style_run(cap_run, size=11, bold=True)
        if rtl:
            set_rtl(cap)


def export_markdown(md_path: Path, docx_path: Path, rtl: bool):
    document = Document()
    apply_doc_defaults(document, rtl)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_code_block = False
    code_buffer: list[str] = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                in_code_block = False
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run("\n".join(code_buffer))
                style_run(run, font_name="Courier New", size=10)
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not stripped:
            document.add_paragraph()
            continue

        if stripped == "---":
            document.add_paragraph("_" * 50)
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            caption, rel_path = image_match.groups()
            add_image(document, (md_path.parent / rel_path).resolve(), caption, rtl)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            hashes, text = heading_match.groups()
            add_heading(document, text, len(hashes), rtl)
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match:
            add_bullet(document, ordered_match.group(1), rtl, ordered=True)
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            add_bullet(document, bullet_match.group(1), rtl, ordered=False)
            continue

        add_paragraph(document, stripped, rtl)

    document.save(docx_path)


def main():
    export_markdown(
        BASE_DIR / "ACADEMIC_REPORT_AR.md",
        BASE_DIR / "ACADEMIC_REPORT_AR.docx",
        rtl=True,
    )
    export_markdown(
        BASE_DIR / "ACADEMIC_REPORT_EN.md",
        BASE_DIR / "ACADEMIC_REPORT_EN.docx",
        rtl=False,
    )
    print("DOCX export completed.")


if __name__ == "__main__":
    main()
